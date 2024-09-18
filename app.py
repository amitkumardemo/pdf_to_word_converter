import streamlit as st
import streamlit.components.v1 as components
from PyPDF2 import PdfReader
from docx import Document
import io
import re

# Function to extract text from a PDF file
def extract_text_from_pdf(pdf_file):
    pdf_reader = PdfReader(pdf_file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() or ""  # Handle cases where extract_text() might return None
    return text

# Main function to run the Streamlit app
def main():
    # Set page configuration
    st.set_page_config(page_title="PDF to Word Converter", layout="wide")

    # Logo
    st.image("jb.png", width=250)  # Replace with your logo URL

    # Navbar
    st.markdown("""
    <nav style="background-color: #007bff; padding: 10px; text-align: center;">
        <a href="#Home" style="color: white; text-decoration: none; padding: 14px 20px;">Home</a>
        <a href="#About" style="color: white; text-decoration: none; padding: 14px 20px;">About</a>
        <a href="https://techiehelpt.netlify.app/" style="color: white; text-decoration: none; padding: 14px 20px;">Back To Website</a>
    </nav>
    """, unsafe_allow_html=True)

    # Main content
    st.title("PDF to Word Converter")
    st.write("Upload a PDF file to extract text and convert it into a Word document.")

    # File uploader for PDF input
    uploaded_file = st.file_uploader("Choose a PDF file...", type="pdf")

    if uploaded_file is not None:
        # Extract text from the PDF
        st.write("Extracting text from PDF...")
        try:
            text = extract_text_from_pdf(uploaded_file)

            # Sanitize text to avoid problematic characters
            text = re.sub(r'[^\x00-\x7F]+', '', text)  # Remove non-ASCII characters
            text = text.replace('\x00', '')  # Remove null bytes
            text = ''.join(ch for ch in text if ch.isprintable())  # Remove non-printable characters

            # Create a Word document
            doc = Document()
            doc.add_heading('Extracted Text', level=1)
            doc.add_paragraph(text)

            # Save the document to a BytesIO stream
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            # Download button for the Word document
            st.download_button(
                label="📥 Download Word Document",
                data=buffer,
                file_name="extracted_text.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

            # Add copy-to-clipboard functionality using components.html
            components.html(f"""
            <script>
            function copyToClipboard() {{
                const text = `{text.replace("`", "\\`")}`;
                navigator.clipboard.writeText(text).then(() => {{
                    alert('Text copied to clipboard!');
                }});
            }}
            </script>
            <button class="icon-btn" onclick="copyToClipboard()">📋 Copy Text</button>
            """, height=50, scrolling=False)

        except Exception as e:
            st.error(f"An error occurred: {e}")

    # Footer
    st.markdown("""
    <footer style="background-color: #007bff; color: white; text-align: center; padding: 20px; margin-top: 20px;">
        <p>© 2024 PDF to Word Converter | TechieHelp</p>
        <a href="https://www.linkedin.com/in/techiehelp" style="color: white; text-decoration: none; margin: 0 10px;">LinkedIn</a>
        <a href="https://www.twitter.com/techiehelp" style="color: white; text-decoration: none; margin: 0 10px;">Twitter</a>
        <a href="https://www.instagram.com/techiehelp2" style="color: white; text-decoration: none;">Instagram</a>
    </footer>
    """, unsafe_allow_html=True)

# Entry point for the Streamlit app
if __name__ == "__main__":
    main()
